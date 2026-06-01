"""DOCX text extractor for CNB OAM documents.

Extracts text from Word documents (.docx) using python-docx.
"""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str | Path) -> str:
    """Extract text content from a Microsoft Word (.docx) document.

    Reads all paragraphs and table contents, concatenating them with newlines.

    Args:
        file_path: Path to the .docx file on disk.

    Returns:
        Extracted text content, or empty string on failure.
    """
    file_path = Path(file_path)
    if not file_path.exists():
        logger.error("DOCX file not found: %s", file_path)
        return ""

    try:
        import docx
        logger.debug("Extracting DOCX text using python-docx: %s", file_path)
        doc = docx.Document(file_path)
        
        text_parts = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in row_text: # simple deduplication of merged cells
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)

    except Exception as e:
        logger.error("Failed to extract text from DOCX %s: %s", file_path.name, e)
        return ""
